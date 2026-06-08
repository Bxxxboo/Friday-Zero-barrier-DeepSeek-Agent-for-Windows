"""生成「星期五 vs Reasonix / DeepSeek-GUI」优化对标 Word 文档到桌面。"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from friday.version import __version__


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Microsoft YaHei"
    run.bold = bold


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val


def build_document() -> Document:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("星期五（Friday）优化对标分析")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.name = "Microsoft YaHei"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"对标 Reasonix · DeepSeek-GUI · DeepSeek-App\n"
        f"版本：Friday v{__version__}　　日期：{date.today().isoformat()}"
    )
    sr.font.size = Pt(11)
    sr.font.name = "Microsoft YaHei"
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    _heading(doc, "一、文档目的", 1)
    _para(
        doc,
        "本文对照 GitHub 上较成熟的 DeepSeek 桌面/终端 Agent（Reasonix、DeepSeek-GUI、"
        "DeepSeek-App 等），梳理星期五当前能力与差距，并给出可落地的优化方向与优先级，"
        "供产品路线与迭代排期参考。",
    )

    _heading(doc, "二、对标项目概览", 1)

    _heading(doc, "2.1 Reasonix（esengine/DeepSeek-Reasonix）", 2)
    _bullet(doc, "定位：DeepSeek 原生终端编码 Agent，约 1.8 万 Star；MIT；Node ≥22。")
    _bullet(doc, "核心：Cache-first Loop（前缀缓存稳定）、Tool-call Repair、R1 思考收割、成本计量。")
    _bullet(doc, "能力：Plan 模式、MCP 原生、Skills/Hooks/Memory、Web 搜索、语义索引、子 Agent 委派。")
    _bullet(doc, "桌面：Tauri 预发布客户端，多 Tab、文件侧栏、Token/缓存命中率仪表。")
    _bullet(doc, "远程：QQ 频道、Discord；配置 ~/.reasonix。")

    _heading(doc, "2.2 DeepSeek-GUI（XingYu-Zhong/DeepSeek-GUI）", 2)
    _bullet(doc, "定位：面向开发者/高频用户的本地工作台，约 2.1k Star；内置 Kun 运行时（TypeScript）。")
    _bullet(doc, "工作台：Code（项目 Agent）、Write（Markdown 写作）、连接手机（飞书/微信/定时）。")
    _bullet(doc, "工程化：新建需求→计划→Todo、/review 代码审查、内联 diff、变更审查面板、/goal 目标模式。")
    _bullet(doc, "扩展：MCP + Skill 图形化管理、mcp_search 按需发现工具、跨会话 Memory、子 Agent。")
    _bullet(doc, "Token ROI：稳定 system prompt、工具结果压缩、DeepSeek 缓存命中可视化。")
    _bullet(doc, "跨平台：macOS / Windows / Linux 安装包；官网 deepseek-gui.com。")

    _heading(doc, "2.3 DeepSeek-App（wzxnb2333/DeepSeek-App）", 2)
    _bullet(doc, "定位：Windows 本地智能体桌面工作台；Electron + Rust sidecar runtime。")
    _bullet(doc, "结构：项目/线程分离；工具审批在输入区附近；思考/工具调用折叠展示。")
    _bullet(doc, "安全：preload 受限 API，renderer 无 Node 权限；随机端口 + 一次性 token。")
    _bullet(doc, "缺口：首版无代码签名与自动更新（SmartScreen 提醒）。")

    _heading(doc, "三、星期五现状摘要（v1.1.3）", 1)
    _bullet(doc, "栈：Python FastAPI + WebView2 无边框壳 + Vanilla JS；Win10/11 专用。")
    _bullet(doc, "优势：中文管家 UI、三级交互模式（Ask/Agent/Yolo）、下载信任链、操作时间线、Gitee 更新源。")
    _bullet(doc, "工具：约 40+ 内置工具（文件/文档/系统/下载/PowerShell/Python/视觉/生图）。")
    _bullet(doc, "差异化：微信 OpenClaw 桥接 + 远程审批；定时任务；配置包迁移；开机自启。")
    _bullet(doc, "Agent：单循环最多 20 轮；工具结果 2KB 截断；无 MCP/子 Agent/Plan 面板。")

    _heading(doc, "四、能力对照矩阵", 1)
    _table(
        doc,
        ["维度", "Reasonix", "DeepSeek-GUI", "星期五", "差距判断"],
        [
            ["目标用户", "开发者/终端", "开发者+写作者", "普通 Windows 用户", "星期五差异化正确"],
            ["Token 成本", "★★★ 缓存优先", "★★★ Kun ROI", "★☆ 未优化缓存", "高优先级"],
            ["MCP 生态", "★★★ 原生", "★★★ GUI 管理", "✗ 无", "中高优先级"],
            ["代码协作", "SEARCH/REPLACE+apply", "diff+review 面板", "文件工具+时间线", "可加强 diff/review"],
            ["Plan/任务分解", "/plan /todo /goal", "需求→计划→Todo", "定时任务+欢迎 chips", "可引入计划面板"],
            ["Memory/RAG", "remember/recall+索引", "跨会话 Memory", "会话内上下文", "中长期"],
            ["IM 远程", "QQ 频道", "飞书/微信/relay", "微信 OpenClaw", "星期五已领先（微信）"],
            ["多平台", "Win/Mac/Linux", "Win/Mac/Linux", "仅 Windows", "战略选择，非必做"],
            ["自动更新", "npm/Tauri", "安装包+更新入口", "Gitee zip 手动", "体验短板"],
            ["安全审批", "Hooks+权限", "模式+审批", "★★★ 成熟", "星期五优势"],
            ["可移植性", "项目 .reasonix", "本地配置", "★★★ zip 迁移", "星期五优势"],
        ],
    )

    _heading(doc, "五、优化建议（按优先级）", 1)

    _heading(doc, "P0 — 短期可落地（1～2 个迭代）", 2)
    _bullet(doc, "Token 经济：稳定 system prompt 与工具 schema 顺序；展示 prompt/completion/cache 用量与估算费用。")
    _bullet(doc, "上下文卫生：超长工具结果智能摘要而非硬截 2KB；重复工具循环检测与早停提示。")
    _bullet(doc, "变更可见性：Agent 改文件后在聊天区或侧栏展示 diff 摘要，一键打开/撤销（借鉴 DeepSeek-GUI review）。")
    _bullet(doc, "更新体验：应用内一键下载并启动更新脚本（或 Squirrel/增量包），减少手动解压 zip。")
    _bullet(doc, "微信/OpenClaw：设置页增加 Gateway 自启开关（与星期五本体自启并列）；失败诊断一键复制。")

    _heading(doc, "P1 — 中期增强（1～2 季度）", 2)
    _bullet(doc, "MCP 客户端：支持 stdio MCP 服务器配置（设置页图形化），按需加载工具（mcp_search 思路）。")
    _bullet(doc, "Plan 模式：长任务生成可编辑计划 Markdown + 线程 Todo，与 Agent 循环联动。")
    _bullet(doc, "Skills 2.0：支持 subagent 模式、Claude 格式 SKILL.md 兼容加载（Reasonix 已验证路径）。")
    _bullet(doc, "Hooks：PreToolUse/PostToolUse 生命周期脚本，满足进阶用户自动化（可选关闭）。")
    _bullet(doc, "定时任务增强：允许白名单内 EXEC（如固定脚本路径）；任务失败通知（系统托盘/微信）。")
    _bullet(doc, "多会话 UX：会话压缩/分叉/归档；旁支对话 /btw 式轻量分支。")

    _heading(doc, "P2 — 战略可选", 2)
    _bullet(doc, "Write 工作台：独立 Markdown 写作空间（星期五已有文档工具，可复用）。")
    _bullet(doc, "语义索引/RAG：工作区文件 embedding 检索，Ask 模式增强。")
    _bullet(doc, "子 Agent 委派：复杂任务 spawn 隔离循环，主会话只收结论。")
    _bullet(doc, "跨平台：评估 Tauri/Electron 壳复用现有 web/，Mac/Linux 次要市场。")
    _bullet(doc, "代码签名：Windows Authenticode，消除 SmartScreen，与 Reasonix/DeepSeek-App 同级信任。")

    _heading(doc, "六、不建议盲目跟进的项", 1)
    _bullet(doc, "做成 IDE 级多面板：星期五定位是「电脑管家」，不是 Cursor/VS Code 替代品。")
    _bullet(doc, "DeepSeek-only 缓存极致优化到 Reasonix 程度：可借鉴机制，但需保留多模型扩展空间。")
    _bullet(doc, "引入 Node/Kun 双运行时：维护成本高；优先 Python 内增强或 MCP 接入外部能力。")
    _bullet(doc, "去掉审批链换「全自动」：星期五目标用户需要信任感，安全仍是核心竞争力。")

    _heading(doc, "七、建议路线图（示意）", 1)
    _table(
        doc,
        ["阶段", "主题", "交付物"],
        [
            ["v1.2", "可观测 + 变更审查", "Token 仪表、diff 摘要、工具结果摘要"],
            ["v1.3", "更新与自启完善", "应用内更新、OpenClaw 自启开关、诊断导出"],
            ["v1.4", "Plan + Todo", "计划面板、线程 Todo、/plan 指令"],
            ["v1.5", "MCP 预览", "stdio MCP 配置、按需工具加载"],
            ["v2.0", "Memory + Hooks", "跨会话记忆、生命周期 Hook（可选）"],
        ],
    )

    _heading(doc, "八、结论", 1)
    _para(
        doc,
        "星期五在「Windows 零门槛、安全审批、微信远程、国内更新、可移植迁移」上已形成清晰差异化，"
        "不宜整体照搬 Reasonix/DeepSeek-GUI 的开发者 IDE 形态。最值得借鉴的是："
        "（1）Token/缓存可观测与上下文卫生；（2）文件变更审查与 Plan/Todo 任务化；"
        "（3）MCP 生态接入；（4）安装更新体验。按 P0→P1 顺序推进，可在保持管家定位的同时"
        "显著拉近与成熟 Agent 工作台的能力差距。",
    )

    _para(doc, "参考链接：", bold=True)
    _bullet(doc, "https://github.com/esengine/DeepSeek-Reasonix")
    _bullet(doc, "https://github.com/XingYu-Zhong/DeepSeek-GUI")
    _bullet(doc, "https://github.com/wzxnb2333/DeepSeek-App")
    _bullet(doc, "https://github.com/deepseek-ai/awesome-deepseek-agent")

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("— 由星期五仓库 scripts/generate_optimization_docx.py 自动生成 —")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


def main() -> None:
    desktop = Path.home() / "Desktop"
    out = desktop / "星期五优化对标分析.docx"
    doc = build_document()
    doc.save(out)
    print(str(out))


if __name__ == "__main__":
    main()
